import os
import time
import threading
import datetime
import tkinter as tk
import logging
from tkinter import filedialog, messagebox, scrolledtext

# 【全局断路器】：强行瘫痪整个 Python 进程中所有 INFO 及以下级别的标准日志输出
logging.disable(logging.INFO)
# 顺手把 PyMuPDF (fitz) 底层的 C++ 警告也物理静音
os.environ["FITZ_LOG"] = "0"

try:
    from pdf2docx import Converter
except ImportError:
    Converter = None


class PDFToWordUI:
    """企业级 PDF 转 Word 智能转换引擎 (防覆盖保护 + 异步扫盘 + 状态机管控)"""

    def __init__(self, parent_frame):
        self.parent = parent_frame

        # 线程控制事件
        self._stop_event = threading.Event()
        self._pause_event = threading.Event()
        self._pause_event.set()

        # UI 变量
        self.source_var = tk.StringVar()
        self.target_var = tk.StringVar()
        self.scan_subdirs_var = tk.BooleanVar(value=True)  # 是否扫描子目录

        # 状态机动画变量
        self.anim_running = False
        self.anim_frame = 0
        self.spinner_frames = ["⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏"]

        self.setup_ui()

    # ================= 1. UI 布局与视图 =================
    def setup_ui(self):
        # 拦截：如果缺少核心库，直接渲染红色警告面板
        if Converter is None:
            tk.Label(self.parent, text="⚠️ 缺少核心运行库！请在终端运行：\npip install pdf2docx",
                     fg="red", font=("Microsoft YaHei UI", 12, "bold"), justify=tk.CENTER).pack(pady=50)
            return

        frame = tk.Frame(self.parent, padx=20, pady=20)
        frame.pack(fill=tk.BOTH, expand=True)
        frame.columnconfigure(1, weight=1)
        frame.rowconfigure(5, weight=1)  # 日志区完美拉伸

        # 1. 数据源区
        tk.Label(frame, text="1. 源 PDF 文件/目录:", font=("Microsoft YaHei UI", 10, "bold")).grid(row=0, column=0, sticky=tk.W, pady=5)
        entry_source = tk.Entry(frame, textvariable=self.source_var, font=("Microsoft YaHei UI", 10))
        entry_source.grid(row=0, column=1, sticky=tk.EW, padx=10, pady=5, ipady=3)
        # 绑定回车事件，粘贴路径后自动联动输出目录
        entry_source.bind("<Return>", lambda e: self._link_target_dir(self.source_var.get()))

        btn_src_frame = tk.Frame(frame)
        btn_src_frame.grid(row=0, column=2, sticky=tk.W)
        tk.Button(btn_src_frame, text="选择文件...", command=self.select_files, width=10).pack(side=tk.LEFT, padx=(0, 5))
        tk.Button(btn_src_frame, text="选择目录...", command=self.select_folder, width=10).pack(side=tk.LEFT)

        # 2. 转换配置区
        tk.Label(frame, text="2. 转换引擎配置:", font=("Microsoft YaHei UI", 10, "bold")).grid(row=1, column=0, sticky=tk.W, pady=5)
        config_frame = tk.Frame(frame)
        config_frame.grid(row=1, column=1, columnspan=2, sticky=tk.W, padx=5, pady=5)
        tk.Checkbutton(config_frame, text="包含所有深层子目录 (仅选择目录时生效)", variable=self.scan_subdirs_var,
                       font=("Microsoft YaHei UI", 9), fg="#2C3E50").pack(side=tk.LEFT)
        tk.Label(config_frame, text="  |  💡 提示：纯扫描件(图片构成的PDF)转换后仍为图片，无法直接修改文字。",
                 font=("Microsoft YaHei UI", 9), fg="#D35400").pack(side=tk.LEFT)

        # 3. 输出目录区
        tk.Label(frame, text="3. Word 保存至:", font=("Microsoft YaHei UI", 10, "bold")).grid(row=2, column=0, sticky=tk.W, pady=5)
        tk.Entry(frame, textvariable=self.target_var, font=("Microsoft YaHei UI", 10)).grid(row=2, column=1, sticky=tk.EW, padx=10, pady=5, ipady=3)
        tk.Button(frame, text="浏览目录...", command=self.select_target, width=12).grid(row=2, column=2, sticky=tk.W, pady=5)

        # 4. 状态机控制面板
        btn_frame = tk.Frame(frame)
        btn_frame.grid(row=3, column=0, columnspan=3, pady=15, sticky=tk.EW)
        btn_frame.columnconfigure((0, 1, 2), weight=1)

        self.btn_start = tk.Button(btn_frame, font=("Microsoft YaHei UI", 11, "bold"))
        self.btn_start.grid(row=0, column=0, padx=5, sticky=tk.EW, ipady=5)
        self.btn_pause = tk.Button(btn_frame, font=("Microsoft YaHei UI", 11, "bold"))
        self.btn_pause.grid(row=0, column=1, padx=5, sticky=tk.EW, ipady=5)
        self.btn_cancel = tk.Button(btn_frame, font=("Microsoft YaHei UI", 11, "bold"))
        self.btn_cancel.grid(row=0, column=2, padx=5, sticky=tk.EW, ipady=5)

        # 5. 富文本日志追踪区
        self.log_area = scrolledtext.ScrolledText(frame, font=("Microsoft YaHei UI", 10), bg="#F8F9FA", fg="#2C3E50",
                                                  relief=tk.FLAT, padx=12, pady=12, state=tk.DISABLED)
        self.log_area.grid(row=5, column=0, columnspan=3, sticky=tk.NSEW, pady=5)

        # 统一规范的标签映射
        self.log_area.tag_config("INFO", foreground="#2C3E50")
        self.log_area.tag_config("SUCCESS", foreground="#198754", font=("Microsoft YaHei UI", 10, "bold"))
        self.log_area.tag_config("WARNING", foreground="#D35400", font=("Microsoft YaHei UI", 10, "bold"))
        self.log_area.tag_config("ERROR", foreground="#DC3545", font=("Microsoft YaHei UI", 10, "bold"))
        self.log_area.tag_config("MAGENTA", foreground="#8E44AD", font=("Microsoft YaHei UI", 10, "bold"))
        self.log_area.tag_config("LINK", foreground="#0078D7", underline=True)
        self.log_area.tag_bind("LINK", "<Enter>", lambda e: self.log_area.config(cursor="hand2"))
        self.log_area.tag_bind("LINK", "<Leave>", lambda e: self.log_area.config(cursor=""))

        self.log_to_ui("✨ 企业级 PDF 智能转换引擎已就绪！", "MAGENTA")
        self.update_button_ui("idle")

    # ================= 2. 状态机与动画核心 =================
    def _animate_btn(self):
        if not self.anim_running: return
        if self._pause_event.is_set():
            frame_char = self.spinner_frames[self.anim_frame % len(self.spinner_frames)]
            self.btn_start.config(text=f"{frame_char} 引擎极速转换中...")
            self.anim_frame += 1
        self.parent.after(100, self._animate_btn)

    def update_button_ui(self, state):
        color_start_active = "#8E44AD"  # 保持与效能工具箱高级功能一致的紫色主题
        color_pause_active = "#F39C12"
        color_stop_active = "#DC3545"
        color_disabled_bg = "#E9ECEF"
        color_disabled_fg = "#ADB5BD"
        color_enabled_gray_bg = "#E0E0E0"
        color_enabled_gray_fg = "#333333"

        if state == "idle":
            self.anim_running = False
            self.btn_start.config(bg=color_start_active, fg="white", text="▶ 开始智能转换", state=tk.NORMAL, command=self.start_conversion_thread)
            self.btn_pause.config(bg=color_disabled_bg, fg=color_disabled_fg, text="⏸ 暂停", state=tk.DISABLED)
            self.btn_cancel.config(bg=color_disabled_bg, fg=color_disabled_fg, text="⏹ 停止", state=tk.DISABLED)
        elif state == "running":
            self.btn_start.config(bg=color_start_active, fg="white", state=tk.NORMAL, command=lambda: None)
            self.btn_pause.config(bg=color_enabled_gray_bg, fg=color_enabled_gray_fg, text="⏸ 暂停", state=tk.NORMAL, command=self.toggle_pause)
            self.btn_cancel.config(bg=color_enabled_gray_bg, fg=color_enabled_gray_fg, text="⏹ 停止", state=tk.NORMAL, command=self.cancel_processing)
            if not self.anim_running:
                self.anim_running = True
                self.anim_frame = 0
                self._animate_btn()
        elif state == "paused":
            self.btn_start.config(bg=color_enabled_gray_bg, fg="#888888", text="⚙️ 已挂起", state=tk.NORMAL, command=lambda: None)
            self.btn_pause.config(bg=color_pause_active, fg="white", text="▶ 继续转换", state=tk.NORMAL, command=self.toggle_pause)
            self.btn_cancel.config(bg=color_enabled_gray_bg, fg=color_enabled_gray_fg, text="⏹ 停止", state=tk.NORMAL, command=self.cancel_processing)
        elif state == "stopping":
            self.anim_running = False
            self.btn_start.config(bg=color_disabled_bg, fg=color_disabled_fg, text="🛑 正在平滑中止...", state=tk.DISABLED)
            self.btn_pause.config(bg=color_disabled_bg, fg=color_disabled_fg, text="⏸ 暂停", state=tk.DISABLED)
            self.btn_cancel.config(bg=color_stop_active, fg="white", text="⏹ 清理缓存...", state=tk.NORMAL, command=lambda: None)

    # ================= 3. 用户交互与联动 =================
    def _link_target_dir(self, source_path):
        """架构师级体验优化：每次更换输入源，输出目录强制无感联动"""
        source_path = source_path.strip().strip('"').strip("'")
        if not source_path or not os.path.exists(source_path):
            return

        if os.path.isfile(source_path):
            base_dir = os.path.dirname(source_path)
        else:
            base_dir = source_path

        # 统一生成专属输出目录
        smart_target_dir = os.path.abspath(os.path.join(base_dir, "Word输出结果"))
        self.target_var.set(smart_target_dir)

    def select_files(self):
        paths = filedialog.askopenfilenames(title="选择需要转换的 PDF", filetypes=[("PDF 文档", "*.pdf")])
        if paths:
            # 存入所有路径，以竖线分割（UI 表现层）
            joined_paths = "|".join(paths)
            self.source_var.set(joined_paths)
            self._link_target_dir(paths[0])

    def select_folder(self):
        path = filedialog.askdirectory(title="选择包含 PDF 的文件夹")
        if path:
            self.source_var.set(os.path.abspath(path))
            self._link_target_dir(path)

    def select_target(self):
        path = filedialog.askdirectory(title="选择保存目录")
        if path: self.target_var.set(os.path.abspath(path))

    def log_to_ui(self, message, level="INFO", link_path=None, auto_scroll=True):
        def update():
            self.log_area.config(state=tk.NORMAL)
            timestamp = datetime.datetime.now().strftime("[%H:%M:%S] ")
            self.log_area.insert(tk.END, f"[{timestamp}] {message}\n", level)
            if link_path:
                tag_name = f"link_{datetime.datetime.now().timestamp()}"
                self.log_area.tag_config(tag_name, foreground="#0078D7", underline=True, font=("Microsoft YaHei UI", 10, "bold"))
                self.log_area.tag_bind(tag_name, "<Enter>", lambda e: self.log_area.config(cursor="hand2"))
                self.log_area.tag_bind(tag_name, "<Leave>", lambda e: self.log_area.config(cursor=""))
                self.log_area.tag_bind(tag_name, "<Button-1>", lambda e, p=link_path: os.startfile(p))
                self.log_area.insert(tk.END, f" ➔ 点击打开查看结果\n", tag_name)
            if auto_scroll: self.log_area.see(tk.END)
            self.log_area.config(state=tk.DISABLED)

        self.parent.after(0, update)

    def toggle_pause(self):
        if self._pause_event.is_set():
            self._pause_event.clear()
            self.update_button_ui("paused")
            self.log_to_ui("⏸ 任务已挂起，等待恢复...", "WARNING")
        else:
            self._pause_event.set()
            self.update_button_ui("running")
            self.log_to_ui("▶ 任务已恢复，继续转换。", "INFO")

    def cancel_processing(self):
        if messagebox.askyesno("确认停止", "确定要中断转换任务吗？\n已转换完成的文件将被保留。"):
            self._stop_event.set()
            self._pause_event.set()  # 防止死锁
            self.update_button_ui("stopping")
            self.log_to_ui("🛑 正在安全中断任务并销毁临时句柄...", "WARNING", auto_scroll=False)

    # ================= 4. 核心转换业务逻辑 =================
    def _get_safe_path(self, target_dir, base_name, ext=".docx"):
        """【防覆盖保护伞】智能防重名与防锁定算法"""
        out_path = os.path.join(target_dir, f"{base_name}{ext}")
        counter = 1
        while os.path.exists(out_path):
            out_path = os.path.join(target_dir, f"{base_name}({counter}){ext}")
            counter += 1
        return out_path

    def start_conversion_thread(self):
        source_raw = self.source_var.get().strip()
        tgt_path = self.target_var.get().strip().strip('"').strip("'")
        scan_subdirs = self.scan_subdirs_var.get()

        if not source_raw or not tgt_path:
            messagebox.showwarning("提示", "请确保源路径和输出目录已填写完整！")
            return

        self._stop_event.clear()
        self._pause_event.set()
        self.update_button_ui("running")

        self.log_area.config(state=tk.NORMAL)
        self.log_area.delete(1.0, tk.END)
        self.log_area.config(state=tk.DISABLED)

        threading.Thread(target=self.process_conversion, args=(source_raw, tgt_path, scan_subdirs), daemon=True).start()

    def _inject_bookmarks_to_word(self, pdf_path, docx_path):
        """架构级大纲重建器 V5.1：无限视距 + 连体剥离 + 智能语义层级校准"""
        try:
            import fitz
            from docx import Document
            import re

            pdf_doc = fitz.open(pdf_path)
            toc = pdf_doc.get_toc()
            pdf_doc.close()

            if not toc: return

            self.log_to_ui("  └─ 💡 侦测到超高复杂度大纲树，启动 V5.1 语义重塑引擎...")

            # ================= 核心升级 3：智能语义层级校准 (Semantic Calibration) =================
            # 解决 PDF 物理结构层级与文本字面语义层级错位的问题 (例如：结构是1级，但文本是 "1.1" 属于2级)
            level_shift = 0
            for lvl, title, _ in toc:
                # 正则精确捕获标准的点分多级编号，例如 "1.1", "1.1.3"
                match = re.match(r'^(\d{1,2}(?:\.\d{1,2})+)(?=[\s\u4e00-\u9fa5A-Za-z]|$)', title.strip())
                if match:
                    number_str = match.group(1)
                    # 通过点号的数量，计算出它的绝对语义层级 (例如 "1.1" 是 2 级)
                    semantic_level = len(number_str.split('.'))

                    # 如果文本字面语义层级 > PDF物理结构层级，计算补偿偏移量
                    if semantic_level > lvl:
                        level_shift = semantic_level - lvl
                        self.log_to_ui(f"  └─ 🔧 触发层级校准：侦测到根节点缺失，全局大纲自动下沉 {level_shift} 级。")
                    break  # 只要找到第一个带编号的基准节点，就校准完毕，跳出循环

            # ===================================================================================

            word_doc = Document(docx_path)

            # ================= 🎨 核心新增：全局中文字体兜底策略 =================
            try:
                from docx.oxml.ns import qn

                # 获取 Word 文档的“正文”基座样式
                normal_style = word_doc.styles['Normal']

                # 1. 设置西文（英文字母/数字）的兜底字体，通常设为 Times New Roman 或与中文统一
                normal_style.font.name = '宋体'

                # 2. 【最关键一步】强行指定“东亚文字（中文）”的兜底字体为宋体
                # 只有修改底层的 w:eastAsia 属性，Word 才会乖乖听话把默认中文渲染成宋体
                normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                self.log_to_ui("  └─ 🎨 已注入全局字体防线：未知中文将平滑回退至 [宋体]。")
            except Exception as e:
                pass  # 防御性编程，字体兜底失败不影响核心流程
            # ======================================================================

            all_paras = []
            for p in word_doc.paragraphs:
                all_paras.append(p)
            for table in word_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            all_paras.append(p)

            def clean_text(text):
                if not text: return ""
                return re.sub(r'[\s\.\-\_、，：:()（）]+', '', str(text)).lower()

            toc_index = 0
            toc_length = len(toc)
            max_lookahead = 200

            for para in all_paras:
                if toc_index >= toc_length:
                    break

                para_text = para.text.strip()
                para_text_clean = clean_text(para_text)
                if not para_text_clean or len(para_text_clean) < 2:
                    continue

                matched_offset = -1
                fused_entries = []

                for offset in range(min(max_lookahead, toc_length - toc_index)):
                    level, title, _ = toc[toc_index + offset]
                    title_clean = clean_text(title)
                    if not title_clean: continue

                    idx = para_text_clean.find(title_clean)
                    if 0 <= idx <= 10:
                        is_physical_toc = False

                        if re.search(r'(\.{3,}|\t+|\s{3,})\d+$', para_text):
                            is_physical_toc = True
                        else:
                            match_para_end = re.search(r'\d+$', para_text)
                            match_title_end = re.search(r'\d+$', title.strip())
                            if match_para_end:
                                if not match_title_end or match_para_end.group() != match_title_end.group():
                                    is_physical_toc = True

                        if not is_physical_toc:
                            temp_offset = offset
                            while (toc_index + temp_offset) < toc_length:
                                t_lvl, t_title, _ = toc[toc_index + temp_offset]
                                t_clean = clean_text(t_title)
                                if t_clean and t_clean in para_text_clean:
                                    fused_entries.append((t_lvl, t_title))
                                    temp_offset += 1
                                else:
                                    break

                            if fused_entries:
                                matched_offset = offset
                                break

                if matched_offset != -1:
                    if len(fused_entries) > 1:
                        for lvl, t_title in fused_entries:
                            try:
                                new_p = para.insert_paragraph_before(t_title)
                                # 应用层级校准偏移量 level_shift
                                new_p.style = f'Heading {min(lvl + level_shift, 9)}'
                            except:
                                pass
                        para.text = ""
                        toc_index += matched_offset + len(fused_entries)

                    else:
                        lvl, t_title = fused_entries[0]
                        title_clean = clean_text(t_title)

                        if len(para_text_clean) <= len(title_clean) + 20:
                            try:
                                # 应用层级校准偏移量 level_shift
                                para.style = f'Heading {min(lvl + level_shift, 9)}'
                            except:
                                pass
                        else:
                            try:
                                new_p = para.insert_paragraph_before(t_title)
                                # 应用层级校准偏移量 level_shift
                                new_p.style = f'Heading {min(lvl + level_shift, 9)}'
                                if t_title in para.text:
                                    para.text = para.text.replace(t_title, "", 1).strip()
                            except:
                                pass

                        toc_index += matched_offset + 1

            word_doc.save(docx_path)
            self.log_to_ui("  └─ ✨ V5.1 重塑成功！已碾碎死锁陷阱并完美校准语义层级。")

        except ImportError:
            pass
        except Exception as e:
            self.log_to_ui(f"  └─ ⚠️ 大纲重建发生局部异常: {str(e)[:50]}", "WARNING")

    def process_conversion(self, source_raw, tgt_path, scan_subdirs):
        start_time = time.time()
        pdf_tasks = []

        try:
            self.log_to_ui("📊 [阶段 1] 正在进行前置深度扫盘...", "MAGENTA")

            # 解析双栖输入源（文件集合 or 单一文件夹）
            if "|" in source_raw:
                pdf_tasks = [p for p in source_raw.split("|") if p.lower().endswith('.pdf') and os.path.exists(p)]
            elif os.path.isfile(source_raw) and source_raw.lower().endswith('.pdf'):
                pdf_tasks.append(source_raw)
            elif os.path.isdir(source_raw):
                for root, dirs, files in os.walk(source_raw):
                    for f in files:
                        if f.lower().endswith('.pdf'):
                            pdf_tasks.append(os.path.join(root, f))
                    if not scan_subdirs:
                        break  # 仅扫描顶层

            if not pdf_tasks:
                self.log_to_ui("⚠️ 未在源路径中发现任何有效的 PDF 文件，任务取消。", "WARNING")
                return

            if not os.path.exists(tgt_path):
                os.makedirs(tgt_path)

            self.log_to_ui(f"  └─ 扫盘完毕，共锁定 {len(pdf_tasks)} 份 PDF 档案。")
            self.log_to_ui(f"\n⚙️ [阶段 2] 启动智能解析引擎...", "MAGENTA")

            success_count = 0
            error_count = 0

            for i, pdf_path in enumerate(pdf_tasks):
                if self._stop_event.is_set():
                    self.log_to_ui(f"❌ 已取消剩余 {len(pdf_tasks) - i} 份文档的转换...", "ERROR", auto_scroll=False)
                    break
                self._pause_event.wait()

                file_name = os.path.basename(pdf_path)
                base_name = os.path.splitext(file_name)[0]

                # 动态计算防冲突保存路径
                safe_out_path = self._get_safe_path(tgt_path, base_name)
                safe_out_filename = os.path.basename(safe_out_path)

                self.log_to_ui(f"[{i + 1}/{len(pdf_tasks)}] 正在抽壳转换: {file_name} ...")

                cv = None
                try:
                    # 启动原生 pdf2docx 解析内核
                    cv = Converter(pdf_path)
                    # 屏蔽库自带的烦人日志，保持我们 UI 的纯净
                    cv.convert(safe_out_path, start=0, end=None)
                    success_count += 1

                    # ⬇️ 【新增挂载点】：原生转换完成后，立刻无缝进行大纲注入重塑
                    self._inject_bookmarks_to_word(pdf_path, safe_out_path)
                    # ⬆️ ========================================================

                except PermissionError:
                    self.log_to_ui(f"  └─ ❌ 写入被拒！Word文件正被其他程序打开: {safe_out_filename}", "ERROR")
                    error_count += 1
                except MemoryError:
                    self.log_to_ui(f"  └─ ❌ 内存爆满！由于该 PDF 体积过大或图片过多导致溢出。", "ERROR")
                    error_count += 1
                except Exception as e:
                    self.log_to_ui(f"  └─ ❌ 转换失败: 文件加密或内部结构损坏 ({str(e)[:30]}...)", "ERROR")
                    error_count += 1
                    # 异常兜底：清理转换失败产生的损坏残骸文件
                    if os.path.exists(safe_out_path):
                        try:
                            os.remove(safe_out_path)
                        except:
                            pass
                finally:
                    # 原生资源销毁
                    if cv:
                        try:
                            cv.close()
                        except:
                            pass
                    # 缓解 UI 拥堵
                    time.sleep(0.01)

            if not self._stop_event.is_set():
                elapsed = time.time() - start_time
                self.log_to_ui(f"\n🎉 批量转换圆满结束！耗时: {elapsed:.2f} 秒。\n战报：成功 {success_count} 份 | 失败 {error_count} 份。", "SUCCESS", link_path=tgt_path)

        except Exception as global_e:
            self.log_to_ui(f"⚠️ 发生不可预知的全局崩溃: {str(global_e)}", "ERROR")
        finally:
            def reset_ui():
                self.update_button_ui("idle")

            self.parent.after(0, reset_ui)